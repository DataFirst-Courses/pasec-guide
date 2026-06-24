#!/usr/bin/env python3
"""Convert Whats in the data.docx to R Markdown with coloured HTML tables."""

from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


MAIN_SECTIONS = {
    "Grade 2 student assessment",
    "Grade 6 student assessment",
    "Student Reading and Mathematics Proficiency Scores",
    "Grade 2 student contextual questionnaire",
    "Grade 6 student contextual questionnaire",
    "Teacher questionnaire",
    "Principal questionnaire",
    "Derived variables",
    "Administrative variables",
}

SUBSECTION_PREFIXES = (
    "Areas assessed by PASEC2019",
    "Reading proficiency scale",
    "Mathematics proficiency scale",
)


def paragraph_text_plain(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def paragraph_all_bold(paragraph: Paragraph) -> bool:
    runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(runs) and all(run.bold for run in runs)


def paragraph_to_markdown(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        escaped = text
        if run.bold and run.italic:
            escaped = f"***{escaped}***"
        elif run.bold:
            escaped = f"**{escaped}**"
        elif run.italic:
            escaped = f"*{escaped}*"
        parts.append(escaped)
    return "".join(parts).strip()


def word_color_to_css(value: str | None) -> str | None:
    if not value or value.lower() == "auto":
        return None
    if value.lower() == "ffffff":
        return "#FFFFFF"
    if re.fullmatch(r"[0-9A-Fa-f]{6}", value):
        return f"#{value.upper()}"
    return None


def cell_fill(cell) -> str | None:
    tc = cell._tc
    if tc.tcPr is None:
        return None
    shading = tc.tcPr.find(qn("w:shd"))
    if shading is None:
        return None
    return word_color_to_css(shading.get(qn("w:fill")))


def cell_text_color(cell) -> str | None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run._element.rPr is None:
                continue
            color = run._element.rPr.find(qn("w:color"))
            if color is None:
                continue
            css = word_color_to_css(color.get(qn("w:val")))
            if css:
                return css
    return None


def cell_style_attr(cell) -> str:
    styles: list[str] = []
    fill = cell_fill(cell)
    color = cell_text_color(cell)
    if fill:
        styles.append(f"background-color: {fill}")
    if color:
        styles.append(f"color: {color}")
    if not styles:
        return ""
    return f' style="{"; ".join(styles)}"'


def grid_span(cell) -> int:
    tc = cell._tc
    if tc.tcPr is None:
        return 1
    span = tc.tcPr.find(qn("w:gridSpan"))
    if span is None:
        return 1
    return int(span.get(qn("w:val"), 1))


def v_merge_state(cell) -> str | None:
    tc = cell._tc
    if tc.tcPr is None:
        return None
    merge = tc.tcPr.find(qn("w:vMerge"))
    if merge is None:
        return None
    return merge.get(qn("w:val"), "continue")


def unique_row_cells(row):
    seen: set[int] = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        yield cell


def table_to_html(table) -> str:
    lines = ['<table class="pasec-data-table">']

    for row in table.rows:
        cells = list(unique_row_cells(row))
        if not cells:
            continue

        lines.append("  <tr>")
        for cell in cells:
            merge = v_merge_state(cell)
            if merge == "continue":
                continue

            colspan = grid_span(cell)
            attrs = cell_style_attr(cell)
            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ""
            text = html.escape(cell.text.replace("\n", " ").strip())
            lines.append(f"    <td{colspan_attr}{attrs}>{text}</td>")
        lines.append("  </tr>")

    lines.append("</table>")
    return "\n".join(lines)


def heading_level(paragraph: Paragraph) -> int | None:
    text = paragraph_text_plain(paragraph)
    if not paragraph_all_bold(paragraph):
        return None
    if text.lower().startswith("what"):
        return 1
    if text in MAIN_SECTIONS:
        return 2
    if any(text.lstrip().startswith(prefix) for prefix in SUBSECTION_PREFIXES):
        return 3
    return 2


def convert_docx(docx_path: Path, rmd_path: Path) -> None:
    document = Document(str(docx_path))
    body = document.element.body
    table_index = 0
    output: list[str] = [
        "# What's in the data? {#ch-whats-in-data}",
        "",
    ]
    title_written = False

    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = Paragraph(child, document)
            text = paragraph_text_plain(paragraph)
            if not text:
                continue

            if text.startswith("Source:"):
                output.append(f"*{text}*")
                output.append("")
                continue

            level = heading_level(paragraph)
            if level == 1 and not title_written:
                title_written = True
                continue
            if level:
                hashes = "#" * level
                output.append(f"{hashes} {text.lstrip()}")
                output.append("")
                continue

            output.append(paragraph_to_markdown(paragraph))
            output.append("")

        elif tag == "tbl":
            table = document.tables[table_index]
            table_index += 1
            output.append('<div class="pasec-table-wrap">')
            output.append("")
            output.append(table_to_html(table))
            output.append("")
            output.append("</div>")
            output.append("")

    css_block = """
```{css, echo=FALSE}
.pasec-table-wrap {
  overflow-x: auto;
  margin: 1.25em 0;
}

.pasec-data-table {
  width: 100%;
  border-collapse: collapse;
  font-size: 0.92em;
}

.pasec-data-table td {
  border: 1px solid #d0d0d0;
  padding: 0.45em 0.65em;
  vertical-align: top;
}

.pasec-data-table td[style*="background-color: #002060"]:not([style*="color:"]) {
  color: #FFFFFF;
}
```
"""

    content = "\n".join(output).rstrip() + "\n"
    final = content.replace(
        "# What's in the data? {#ch-whats-in-data}\n\n",
        "# What's in the data? {#ch-whats-in-data}\n" + css_block + "\n",
        1,
    )
    final = final.replace(
        "CLICK HERE TO LEARN MORE ABOUT PLAUSIBLE VALUES AND HOW TO USE THEM IN YOUR ANALYSIS",
        "[Learn more about plausible values and how to use them in your analysis](c03-plausible-values.html)",
    )
    final = final.replace(
        "CLICK HERE FOR MORE INFORMATION ON REPLICATE WEIGHTS AND HOW TO USE THEM",
        "[Learn more about replicate weights and how to use them](c02-replicate-weights.html)",
    )
    rmd_path.write_text(final, encoding="utf-8")


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    docx_path = root / "Whats in the data.docx"
    rmd_path = root / "01-whats-in-the-data.Rmd"
    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
    if len(sys.argv) > 2:
        rmd_path = Path(sys.argv[2])

    convert_docx(docx_path, rmd_path)
    print(f"Wrote {rmd_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
